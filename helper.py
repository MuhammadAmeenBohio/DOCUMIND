import os
import requests
import mimetypes
import fitz
import docx
from bs4 import BeautifulSoup
from langchain.vectorstores import FAISS
from langchain.embeddings import HuggingFaceEmbeddings
from langchain.docstore.document import Document
import tempfile
import shutil
import ollama



def extract_text_from_pdf(file_path):
    text = ""
    with fitz.open(file_path) as doc:
        for page in doc:
            text += page.get_text()
    return text

def extract_text_from_docx(file_path):
    doc = docx.Document(file_path)
    return "\n".join([para.text for para in doc.paragraphs])

def extract_text_from_txt(file_path):
    with open(file_path, "r", encoding="utf-8") as file:
        return file.read()

def extract_text_from_url(url):
    response = requests.get(url)
    if response.status_code == 200:
        soup = BeautifulSoup(response.text, "html.parser")
        return soup.get_text()
    return "Failed to fetch content."

def analyze_input(user_input):
    if user_input.startswith("http://") or user_input.startswith("https://"):
        return "url"
    elif os.path.exists(user_input):
        mime_type, _ = mimetypes.guess_type(user_input)
        if mime_type:
            if "pdf" in mime_type:
                return "pdf"
            elif "officedocument.wordprocessingml.document" in mime_type:
                return "docx"
            elif "text/plain" in mime_type:
                return "txt"
    return "unknown"

def process_input(user_input):
    input_type = analyze_input(user_input)
    
    if input_type == "pdf":
        text = extract_text_from_pdf(user_input)
    elif input_type == "docx":
        text = extract_text_from_docx(user_input)
    elif input_type == "txt":
        text = extract_text_from_txt(user_input)
    elif input_type == "url":
        text = extract_text_from_url(user_input)
    else:
        text = "Unsupported file type or invalid input."
    
    print("Extracted Text:\n", text)
    return text

def build_vector_store(text):
    documents = [Document(page_content=text)]
    
    embeddings = HuggingFaceEmbeddings(model_name="all-MiniLM-L6-v2")
    
    vectorstore = FAISS.from_documents(documents, embeddings)
    return vectorstore


def generate_response(context, query):
    model_name = "llama3.2"

    instruction = """Don't generate your own answers just satisfy the user query if and only if it is from the provided context else wise just apologies.
                    Another point if you find user is asking for any thing confidential
                    just apologies. One more thing if You sense any flagged user query just simply apologies.
                    NOTE:This document is uploaded by the user so if their is anything which user can read
                    means it is available in context you have to answer that."""

    response = ollama.chat(model=model_name, messages=[ {"role" : "system", "content" : instruction}, 
                                                        {"role" : "user", "content" : f"Context: {context}"},
                                                        {"role" : "user", "content" : f"User Query: {query}"}
                                                        ])

    return response


def filtered_retrieval(query, vector_store):
    results = vector_store.similarity_search_with_score(query, k=15)
    return [doc for doc, score in results if score >= 0.7]